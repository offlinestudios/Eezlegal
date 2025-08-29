from flask import Blueprint, jsonify, request, current_app, send_file
import os
import uuid
from werkzeug.utils import secure_filename
import PyPDF2
from docx import Document as DocxDocument
import openai
from src.models.user import User, Document, db
from src.routes.user import token_required

documents_bp = Blueprint('documents', __name__)

# Allowed file extensions
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'doc', 'txt'}

def allowed_file(filename):
    """Check if file extension is allowed"""
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_pdf(file_path):
    """Extract text from PDF file"""
    try:
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text.strip()
    except Exception as e:
        print(f"Error extracting PDF text: {str(e)}")
        return None

def extract_text_from_docx(file_path):
    """Extract text from DOCX file"""
    try:
        doc = DocxDocument(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text.strip()
    except Exception as e:
        print(f"Error extracting DOCX text: {str(e)}")
        return None

def extract_text_from_txt(file_path):
    """Extract text from TXT file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read().strip()
    except Exception as e:
        print(f"Error extracting TXT text: {str(e)}")
        return None

def extract_text_from_file(file_path, content_type):
    """Extract text based on file type"""
    if 'pdf' in content_type:
        return extract_text_from_pdf(file_path)
    elif 'word' in content_type or 'docx' in content_type:
        return extract_text_from_docx(file_path)
    elif 'text' in content_type:
        return extract_text_from_txt(file_path)
    else:
        return None

def analyze_document_with_ai(text, document_type="legal document"):
    """Analyze document using OpenAI API"""
    try:
        prompt = f"""Analyze this {document_type} and provide a comprehensive summary including:

1. **Document Type & Purpose**: What kind of document this is and its main purpose
2. **Key Terms & Clauses**: Important provisions, obligations, and rights
3. **Potential Risks**: Areas of concern or unfavorable terms
4. **Recommendations**: Suggested actions or areas to negotiate
5. **Plain English Summary**: A simple explanation of what this document means

Document text:
{text[:4000]}  # Limit to first 4000 characters

Please provide a clear, structured analysis that helps the user understand this document."""

        response = openai.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are a legal document analysis expert. Provide clear, practical analysis of legal documents in plain English."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=1000,
            temperature=0.3
        )
        
        return response.choices[0].message.content
        
    except Exception as e:
        print(f"AI Analysis Error: {str(e)}")
        return "Unable to analyze document at this time. Please try again later."

@documents_bp.route('/documents', methods=['GET'])
@token_required
def get_documents(current_user):
    """Get user's uploaded documents"""
    documents = Document.query.filter_by(user_id=current_user.id)\
                             .order_by(Document.created_at.desc()).all()
    
    return jsonify([doc.to_dict() for doc in documents]), 200

@documents_bp.route('/documents/upload', methods=['POST'])
@token_required
def upload_document(current_user):
    """Upload and process a document"""
    try:
        # Check if file is present
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
        
        file = request.files['file']
        
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        if not allowed_file(file.filename):
            return jsonify({'error': 'File type not allowed. Please upload PDF, DOCX, DOC, or TXT files.'}), 400
        
        # Generate unique filename
        original_filename = secure_filename(file.filename)
        file_extension = original_filename.rsplit('.', 1)[1].lower()
        unique_filename = f"{uuid.uuid4()}.{file_extension}"
        
        # Save file
        file_path = os.path.join(current_app.config['UPLOAD_FOLDER'], unique_filename)
        file.save(file_path)
        
        # Get file info
        file_size = os.path.getsize(file_path)
        content_type = file.content_type or 'application/octet-stream'
        
        # Create document record
        document = Document(
            user_id=current_user.id,
            filename=unique_filename,
            original_filename=original_filename,
            file_path=file_path,
            file_size=file_size,
            content_type=content_type,
            analysis_status='pending'
        )
        
        db.session.add(document)
        db.session.commit()
        
        # Extract text in background (for now, do it synchronously)
        try:
            extracted_text = extract_text_from_file(file_path, content_type)
            
            if extracted_text:
                document.extracted_text = extracted_text
                document.analysis_status = 'processing'
                db.session.commit()
                
                # Analyze with AI
                analysis = analyze_document_with_ai(extracted_text)
                document.analysis_summary = analysis
                document.is_analyzed = True
                document.analysis_status = 'completed'
            else:
                document.analysis_status = 'failed'
            
            db.session.commit()
            
        except Exception as e:
            print(f"Error processing document: {str(e)}")
            document.analysis_status = 'failed'
            db.session.commit()
        
        return jsonify(document.to_dict()), 201
        
    except Exception as e:
        db.session.rollback()
        print(f"Upload error: {str(e)}")
        return jsonify({'error': 'Failed to upload document'}), 500

@documents_bp.route('/documents/<int:document_id>', methods=['GET'])
@token_required
def get_document(current_user, document_id):
    """Get document details"""
    document = Document.query.filter_by(
        id=document_id,
        user_id=current_user.id
    ).first_or_404()
    
    return jsonify({
        'document': document.to_dict(),
        'extracted_text': document.extracted_text,
        'analysis_summary': document.analysis_summary
    }), 200

@documents_bp.route('/documents/<int:document_id>/download', methods=['GET'])
@token_required
def download_document(current_user, document_id):
    """Download original document"""
    document = Document.query.filter_by(
        id=document_id,
        user_id=current_user.id
    ).first_or_404()
    
    if not os.path.exists(document.file_path):
        return jsonify({'error': 'File not found'}), 404
    
    return send_file(
        document.file_path,
        as_attachment=True,
        download_name=document.original_filename
    )

@documents_bp.route('/documents/<int:document_id>/analyze', methods=['POST'])
@token_required
def reanalyze_document(current_user, document_id):
    """Re-analyze document with custom prompt"""
    try:
        document = Document.query.filter_by(
            id=document_id,
            user_id=current_user.id
        ).first_or_404()
        
        if not document.extracted_text:
            return jsonify({'error': 'No text available for analysis'}), 400
        
        data = request.json
        custom_prompt = data.get('prompt', '')
        
        if custom_prompt:
            # Custom analysis with user prompt
            full_prompt = f"""Based on this legal document, please answer the following question:

Question: {custom_prompt}

Document text:
{document.extracted_text[:4000]}

Please provide a detailed, helpful response."""
            
            response = openai.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "You are a legal expert providing analysis of legal documents."},
                    {"role": "user", "content": full_prompt}
                ],
                max_tokens=800,
                temperature=0.3
            )
            
            analysis_result = response.choices[0].message.content
        else:
            # Re-run standard analysis
            analysis_result = analyze_document_with_ai(document.extracted_text)
            document.analysis_summary = analysis_result
            document.is_analyzed = True
            document.analysis_status = 'completed'
            db.session.commit()
        
        return jsonify({
            'analysis': analysis_result,
            'document': document.to_dict()
        }), 200
        
    except Exception as e:
        print(f"Analysis error: {str(e)}")
        return jsonify({'error': 'Failed to analyze document'}), 500

@documents_bp.route('/documents/<int:document_id>', methods=['DELETE'])
@token_required
def delete_document(current_user, document_id):
    """Delete document"""
    try:
        document = Document.query.filter_by(
            id=document_id,
            user_id=current_user.id
        ).first_or_404()
        
        # Delete file from filesystem
        if os.path.exists(document.file_path):
            os.remove(document.file_path)
        
        # Delete from database
        db.session.delete(document)
        db.session.commit()
        
        return jsonify({'message': 'Document deleted successfully'}), 200
        
    except Exception as e:
        db.session.rollback()
        print(f"Delete error: {str(e)}")
        return jsonify({'error': 'Failed to delete document'}), 500

